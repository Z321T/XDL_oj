import torch
import json
import subprocess
from docx import Document
from django.shortcuts import get_object_or_404
from transformers import AutoTokenizer, AutoModel
from torch.nn.functional import cosine_similarity
#cppcheck
from django.db.models.signals import post_save
from django.dispatch import receiver

from CodeBERT_app.models import (CodeFeature, ProgrammingCodeFeature, ProgrammingReportFeature,
                                 ReportStandardScore, CodeStandardScore)
from administrator_app.models import ProgrammingExercise
from teacher_app.models import ExerciseQuestion, ExamQuestion, ReportScore, Teacher

# 加载 CodeBERT 模型和分词器
tokenizer = AutoTokenizer.from_pretrained("microsoft/codebert-base")
model = AutoModel.from_pretrained("microsoft/codebert-base")


# 分析练习与考试代码
def analyze_code(student, code, types, question_id=None):
    # 分词
    tokenized_code = tokenizer.tokenize(code)
    features = []
    for i in range(0, len(tokenized_code), 512):
        inputs = tokenizer(tokenized_code[i:i+512], return_tensors="pt", padding=True, truncation=True, max_length=512)
        # 获取特征值
        with torch.no_grad():
            feature = model(**inputs).last_hidden_state.mean(dim=1)
            features.append(feature)
    # 连接所有特征值
    concatenated_features = torch.cat(features, dim=0)
    feature_as_json = json.dumps(concatenated_features.tolist())

    if types == 'exercise':
        question = ExerciseQuestion.objects.get(id=question_id)
        CodeFeature.objects.create(student=student, exercise_question=question, feature=feature_as_json)
    elif types == 'exam':
        question = ExamQuestion.objects.get(id=question_id)
        CodeFeature.objects.create(student=student, exam_question=question, feature=feature_as_json)


# 分析程序设计题代码
def analyze_programming_code(student, code, question_id):
    # 分词
    tokenized_code = tokenizer.tokenize(code)
    features = []
    for i in range(0, len(tokenized_code), 512):
        inputs = tokenizer(tokenized_code[i:i+512], return_tensors="pt", padding=True, truncation=True, max_length=512)
        # 获取特征值
        with torch.no_grad():
            feature = model(**inputs).last_hidden_state.mean(dim=1)
            features.append(feature)
    # 连接所有特征值
    concatenated_features = torch.cat(features, dim=0)
    feature_as_json = json.dumps(concatenated_features.tolist())

    question = ProgrammingExercise.objects.get(id=question_id)
    ProgrammingCodeFeature.objects.create(student=student, programming_question=question, feature=feature_as_json)


# 分析程序设计题报告
def analyze_programming_report(student, report, question_id):
    # 分词
    tokenized_report = tokenizer.tokenize(report)
    features = []
    for i in range(0, len(tokenized_report), 512):
        inputs = tokenizer(tokenized_report[i:i+512], return_tensors="pt", padding=True, truncation=True, max_length=512)
        # 获取特征值
        with torch.no_grad():
            feature = model(**inputs).last_hidden_state.mean(dim=1)
            features.append(feature)
    # 连接所有特征值
    concatenated_features = torch.cat(features, dim=0)
    feature_as_json = json.dumps(concatenated_features.tolist())

    question = ProgrammingExercise.objects.get(id=question_id)
    ProgrammingReportFeature.objects.create(student=student, programming_question=question, feature=feature_as_json)


# 计算程序设计报告&代码的相似度
def compute_cosine_similarity(feature_json1, feature_json2):
    feature1 = torch.tensor(json.loads(feature_json1)).float()
    feature2 = torch.tensor(json.loads(feature_json2)).float()

    feature1 = feature1.view(1, -1)
    feature2 = feature2.view(1, -1)

    similarity_tensor = cosine_similarity(feature1, feature2)

    similarity_score = similarity_tensor.item()

    return similarity_score


# 报告规范性评分
def score_report(student, file, programmingexercise_id):
    class_assigned = student.class_assigned
    teacher = get_object_or_404(Teacher, classes_assigned=class_assigned)
    report_scores = teacher.report_scores.all()
    for score in report_scores:
        total_score = score.totalscore
        content_score = score.contents
        firstrow_score = score.firstrow
        fontsize_score = score.fontsize
        image_score = score.image
        pagenum_score = score.pagenum

    doc = Document(file)
    # 1. 检查文档是否含有目录
    for para in doc.paragraphs:
        if "目录" in para.text:
            break
        else:
            total_score += content_score

    # 2. 检查每个段落的首行缩进
    first_line = False
    for para in doc.paragraphs:
        if not para.paragraph_format.first_line_indent or para.paragraph_format.first_line_indent.pt != 36:
            first_line = True
            break
    if first_line:
        total_score += firstrow_score

    # 3. 检查段落中的文字字体大小是否为12
    font_size = False
    for para in doc.paragraphs:
        if font_size:
            break
        for run in para.runs:
            if run.font.size and abs(run.font.size.pt - 12) > 1:
                font_size = True
                break
    if font_size:
        total_score += fontsize_score

    # 4. 检查文档中的图像是否居中对齐
    image_alignment = False
    for shape in doc.inline_shapes:
        if shape.type == 3:
            if shape.text_frame.paragraph.alignment != 1:
                image_alignment = True
                break
    if image_alignment:
        total_score += image_score

    # 5. 计算有页码的页面占比
    def score_page_numbers(doc):
        total_pages = len(doc.sections)
        for section in doc.sections:
            page_number_found = False
            for para in section.footer.paragraphs:
                if any(char.isdigit() for char in para.text):
                    page_number_found = True
                    break
            if not page_number_found:
                return True
        if total_pages == 0:
            return True
        return False
    if score_page_numbers(doc):
        total_score += pagenum_score

    # 保存分数
    ReportStandardScore.objects.create(student=student, programming_question_id=programmingexercise_id, standard_score=total_score)


@receiver(post_save, sender=CodeStandardScore)
def run_cppcheck(sender, **kwargs):
    subprocess.call(['cppcheck', '<your options here>'])


# 代码规范性打分
def code_score(cppcheck_output):
    error_scores = {
        'error': 10,
        'warning': 5,
        'performance': 3,
        'style': 1,
    }

    total_score = 0
    for line in cppcheck_output.splitlines():
        for error_type, score in error_scores.items():
            if error_type in line:
                total_score += score

    return total_score
